#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成安徽中烟数字化研发项目投标文件Word版本
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import os

def add_heading_with_style(doc, text, level):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_with_style(doc, text, style='Normal'):
    """添加带样式的段落"""
    para = doc.add_paragraph(text, style=style)
    return para

def parse_markdown_to_word(md_file_path, output_docx_path):
    """将Markdown文件转换为Word文档"""
    
    # 创建Word文档
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)
    
    # 读取Markdown文件
    content = ""
    # 强制读取各个章节文件，确保内容完整
    chapters = [
        'docs/BID_DOCUMENT_MASTER.md',
        'docs/BID_Chapter1_Requirement.md',
        'docs/BID_Chapter2_BizArchitecture.md',
        'docs/BID_Chapter3_BizProcess.md',
        'docs/BID_Chapter4_FuncArchitecture.md',
        'docs/BID_Chapter5_SystemFunctions.md',
        'docs/BID_Chapter6_Technical.md',
        'docs/BID_Chapter7_Implementation.md',
        'docs/BID_Chapter8_Service.md'
    ]
    
    print("开始读取章节文件...")
    for chapter in chapters:
        try:
            with open(chapter, 'r', encoding='utf-8') as f:
                chapter_content = f.read()
                print(f"读取 {chapter}: {len(chapter_content)} 字符")
                content += chapter_content + "\n\n"
        except Exception as e:
            print(f"Error reading {chapter}: {e}")
            
    print(f"总内容长度: {len(content)} 字符")

    
    # 添加封面
    title = doc.add_heading('安徽中烟工业有限责任公司技术中心\n实施数字化研发项目', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('投标文件', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run('项目名称：安徽中烟工业有限责任公司技术中心实施数字化研发项目\n').font.size = Pt(14)
    info_para.add_run('日期：2026年1月\n').font.size = Pt(14)
    
    # 添加分页符
    doc.add_page_break()
    
    # 解析Markdown内容
    lines = content.split('\n')
    in_code_block = False
    in_table = False
    table_lines = []
    
    code_lang = ''
    code_content = []
    
    # 正则：匹配括号内的英文（允许少量标点和数字）
    english_pattern = re.compile(r'\s*\([A-Za-z0-9\s\-_&/,\.]+\)')
    
    # 字体设置辅助函数
    def set_font_style(run, size=Pt(10.5)):
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = size
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.italic = False
    
    def process_table(doc, table_lines):
        """处理Markdown表格并插入Word"""
        # 过滤分隔行 |---|
        rows = [line for line in table_lines if '---' not in line]
        if not rows:
            return

        # 获取列数
        cols = len(rows[0].strip().strip('|').split('|'))
        
        table = doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Table Grid' # 加上边框
        
        for r_idx, row_text in enumerate(rows):
            # 清洗并分割单元格
            cells_text = row_text.strip().strip('|').split('|')
            row_cells = table.rows[r_idx].cells
            
            for c_idx, cell_text in enumerate(cells_text):
                if c_idx < len(row_cells):
                    # 清洗文本
                    clean_text = cell_text.replace('**', '').strip()
                    clean_text = english_pattern.sub('', clean_text)
                    
                    cell = row_cells[c_idx]
                    # 清除默认段落
                    cell._element.clear_content()
                    para = cell.add_paragraph(clean_text)
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.line_spacing = 1.0 # 表格内行距稍微紧凑点
                    if para.runs:
                         set_font_style(para.runs[0])
                    else:
                         # 确保空单元格也有格式（比如后面追加内容时）
                         pass
                    
                    # 重新遍历run设置样式（add_paragraph可能没有run如果text为空）
                    if clean_text:
                         for run in para.runs:
                             set_font_style(run)

    for line in lines:
        stripped_line = line.strip()
        
        # 1. 优先处理代码块状态
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lang = line[3:].strip()
                code_content = []
                # 如果之前在处理表格，需要结束表格
                if in_table:
                    process_table(doc, table_lines)
                    in_table = False
                    table_lines = []
            else:
                in_code_block = False
                # 忽略Mermaid，其他当作普通文本引用
                if code_lang != 'mermaid':
                     para = doc.add_paragraph('\n'.join(code_content), style='Intense Quote')
                code_content = []
            continue
        
        if in_code_block:
            code_content.append(line)
            continue

        # 2. 处理表格状态
        if stripped_line.startswith('|') and stripped_line.endswith('|'): # 简单的Markdown表格检测
            if not in_table:
                # 检查是否真的是表格（看是否多列）
                if len(stripped_line.split('|')) > 2:
                    in_table = True
                    table_lines = [stripped_line]
                else:
                    # 可能是为了样式的竖线，暂不处理为表格，当作普通文本
                    pass 
            else:
                table_lines.append(stripped_line)
            continue
        else:
            # 遇到非表格行，检查是否需要提交表格
            if in_table:
                process_table(doc, table_lines)
                in_table = False
                table_lines = []

        # 3. 优先处理图片
        if stripped_line.startswith('!['):
            match = re.match(r'!\[(.*?)\]\((.*?)\)', stripped_line)
            if match:
                alt = match.group(1)
                img_path = match.group(2)
                try:
                    real_path = img_path
                    if not os.path.exists(real_path):
                         if os.path.exists('docs/' + real_path):
                             real_path = 'docs/' + real_path
                         elif os.path.exists('docs/ppt/' + os.path.basename(real_path)):
                             real_path = 'docs/ppt/' + os.path.basename(real_path)
                    
                    if os.path.exists(real_path):
                        doc.add_picture(real_path, width=Inches(6.0))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if alt:
                            clean_alt = alt.replace('**', '')
                            clean_alt = english_pattern.sub('', clean_alt)
                            caption = doc.add_paragraph(clean_alt, style='Caption')
                            caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption.paragraph_format.line_spacing = 1.5
                            for run in caption.runs:
                                set_font_style(run)
                except Exception as e:
                    print(f"Error adding image {img_path}: {e}")
            continue

        # 4. 文本清洗
        clean_line = line.replace('**', '')
        clean_line = english_pattern.sub('', clean_line)
        
        # 5. 处理标题
        if clean_line.startswith('#'):
            level = len(re.match(r'^#+', clean_line).group())
            text = clean_line.lstrip('#').strip()
            if text:
                heading = add_heading_with_style(doc, text, level)
                for run in heading.runs:
                    set_font_style(run, size=Pt(16 - level)) # 标题大小简单递减
                    run.font.bold = True # 标题保持加粗
        
        # 6. 处理列表
        elif clean_line.strip().startswith('*') or clean_line.strip().startswith('-'):
            text = re.sub(r'^[\s\*\-]+', '', clean_line).strip()
            if text:
                # 不缩进列表
                para = doc.add_paragraph(text, style='List Bullet')
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    set_font_style(run)
                    
        # 7. 处理普通段落 (首行缩进)
        elif clean_line.strip():
            if clean_line.strip() == '---': continue
            if clean_line.strip().startswith('>'): clean_line = clean_line.lstrip('>').strip()
            
            para = doc.add_paragraph(clean_line.strip())
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = Pt(21) # 2字符缩进 (10.5pt * 2)
            
            for run in para.runs:
                set_font_style(run)
    
    # 循环结束，检查是否有遗留表格
    if in_table:
         process_table(doc, table_lines)

    # 保存文档
    doc.save(output_docx_path)
    print(f"Word文档已生成: {output_docx_path}")

if __name__ == '__main__':
    md_file = 'docs/BID_DOCUMENT_MASTER.md' 
    output_file = '安徽中烟数字化研发项目投标文件.docx'
    parse_markdown_to_word(md_file, output_file)

